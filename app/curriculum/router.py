"""
Curriculum upload API router.
Provides endpoints for parsing and importing curriculum files (XLSX/DOCX).
"""

from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import List, Optional
import io

from app.auth.dependencies import get_current_coordinator_id

router = APIRouter(prefix="/api/curriculum", tags=["curriculum"])


class ParsedSubject(BaseModel):
    course_code: str
    course_name: str
    l: int
    t: int
    p: int
    credits: int
    course_category: str
    program_name: str
    semester_label: str
    section_label: str
    shift: int
    student_strength: int
    curriculum_year: str


class ParseResponse(BaseModel):
    success: bool
    subjects: List[ParsedSubject]
    message: str
    warnings: List[str] = []


class ConfirmRequest(BaseModel):
    subjects: List[ParsedSubject]


@router.post("/parse", response_model=ParseResponse)
async def parse_curriculum_file(
    file: UploadFile = File(...),
    coordinator_id: int = Depends(get_current_coordinator_id)
):
    """
    Parse uploaded curriculum file (XLSX or DOCX) and extract subject data.
    Returns parsed subjects for preview before confirmation.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = file.filename.lower().split('.')[-1]
    if file_ext not in ['xlsx', 'xls', 'docx', 'doc']:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload XLSX or DOCX file."
        )
    
    try:
        content = await file.read()
        
        if file_ext in ['xlsx', 'xls']:
            subjects = await parse_excel(content)
        else:
            subjects = await parse_docx(content)
        
        return ParseResponse(
            success=True,
            subjects=subjects,
            message=f"Successfully parsed {len(subjects)} subjects from {file.filename}",
            warnings=[]
        )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse file: {str(e)}"
        )


@router.post("/confirm")
async def confirm_curriculum_import(
    data: ConfirmRequest,
    coordinator_id: int = Depends(get_current_coordinator_id)
):
    """
    Confirm and import parsed subjects into the database.
    Creates subject offerings for all subjects in the request.
    """
    from app.db.session import get_transaction
    from app.subjects import service
    
    imported_count = 0
    failed_count = 0
    errors = []
    
    with get_transaction() as session:
        for subject in data.subjects:
            try:
                # Resolve program, semester, section IDs from names/labels
                program_id = resolve_program_id(session, subject.program_name)
                semester_id = resolve_semester_id(session, subject.semester_label)
                section_id = resolve_section_id(session, subject.section_label, subject.shift)
                
                if not program_id or not semester_id or not section_id:
                    failed_count += 1
                    errors.append(f"{subject.course_code}: Missing program/semester/section mapping")
                    continue
                
                # Create offering
                offering_data = {
                    "course_code": subject.course_code,
                    "course_name": subject.course_name,
                    "l": subject.l,
                    "t": subject.t,
                    "p": subject.p,
                    "credits": subject.credits,
                    "course_category": subject.course_category,
                    "program_id": program_id,
                    "semester_id": semester_id,
                    "section_id": section_id,
                    "shift": subject.shift,
                    "student_strength": subject.student_strength,
                    "curriculum_year": subject.curriculum_year
                }
                
                result = service.create_offering(session, offering_data)
                if result["success"]:
                    imported_count += 1
                else:
                    failed_count += 1
                    errors.append(f"{subject.course_code}: {result['message']}")
            
            except Exception as e:
                failed_count += 1
                errors.append(f"{subject.course_code}: {str(e)}")
        
        session.commit()
    
    return {
        "success": True,
        "imported": imported_count,
        "failed": failed_count,
        "errors": errors,
        "message": f"Import complete: {imported_count} subjects imported, {failed_count} failed"
    }


async def parse_excel(content: bytes) -> List[ParsedSubject]:
    """Parse Excel file and extract subject data."""
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="openpyxl not installed. Please install it to parse Excel files."
        )
    
    workbook = openpyxl.load_workbook(io.BytesIO(content))
    sheet = workbook.active
    
    subjects = []
    
    # Skip header row, start from row 2
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row[0]:  # Skip empty rows
            continue
        
        try:
            subject = ParsedSubject(
                course_code=str(row[0]).strip(),
                course_name=str(row[1]).strip(),
                l=int(row[2] or 0),
                t=int(row[3] or 0),
                p=int(row[4] or 0),
                credits=int(row[5] or 0),
                course_category=str(row[6] or 'CC').strip(),
                program_name=str(row[7]).strip(),
                semester_label=str(row[8]).strip(),
                section_label=str(row[9]).strip(),
                shift=int(row[10] or 1),
                student_strength=int(row[11] or 0),
                curriculum_year=str(row[12] or '2022').strip()
            )
            subjects.append(subject)
        except (IndexError, ValueError, TypeError) as e:
            # Skip malformed rows
            continue
    
    return subjects


async def parse_docx(content: bytes) -> List[ParsedSubject]:
    """Parse DOCX file and extract subject data from tables."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Please install it to parse DOCX files."
        )
    
    doc = Document(io.BytesIO(content))
    subjects = []
    
    # Find first table in document
    if not doc.tables:
        raise HTTPException(status_code=400, detail="No tables found in DOCX file")
    
    table = doc.tables[0]
    
    # Skip header row, start from row 1
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) < 13:
            continue
        
        try:
            subject = ParsedSubject(
                course_code=cells[0].text.strip(),
                course_name=cells[1].text.strip(),
                l=int(cells[2].text.strip() or 0),
                t=int(cells[3].text.strip() or 0),
                p=int(cells[4].text.strip() or 0),
                credits=int(cells[5].text.strip() or 0),
                course_category=cells[6].text.strip() or 'CC',
                program_name=cells[7].text.strip(),
                semester_label=cells[8].text.strip(),
                section_label=cells[9].text.strip(),
                shift=int(cells[10].text.strip() or 1),
                student_strength=int(cells[11].text.strip() or 0),
                curriculum_year=cells[12].text.strip() or '2022'
            )
            subjects.append(subject)
        except (ValueError, TypeError, AttributeError):
            # Skip malformed rows
            continue
    
    return subjects


def resolve_program_id(session, program_name: str) -> Optional[int]:
    """Resolve program name to program ID."""
    from sqlalchemy import text
    result = session.execute(
        text("SELECT id FROM programs WHERE name = :name"),
        {"name": program_name}
    ).fetchone()
    return result[0] if result else None


def resolve_semester_id(session, semester_label: str) -> Optional[int]:
    """Resolve semester label to semester ID."""
    from sqlalchemy import text
    result = session.execute(
        text("SELECT id FROM semesters WHERE label = :label"),
        {"label": semester_label}
    ).fetchone()
    return result[0] if result else None


def resolve_section_id(session, section_label: str, shift: int) -> Optional[int]:
    """Resolve section label and shift to section ID."""
    from sqlalchemy import text
    result = session.execute(
        text("SELECT id FROM sections WHERE label = :label AND shift = :shift"),
        {"label": section_label, "shift": shift}
    ).fetchone()
    return result[0] if result else None
